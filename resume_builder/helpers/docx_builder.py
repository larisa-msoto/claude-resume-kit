#!/usr/bin/env python3
"""
Build a resume or cover letter .docx from a content YAML file.

Usage:
  python3 docx_builder.py render resume resume_content.yaml output/Acme/e2e_acme_resume.docx
  python3 docx_builder.py render cover_letter cl_content.yaml output/Acme/e2e_acme_cover_letter.docx

Style constants (fonts, margins, colors) live in STYLE below -- change them
in one place to retune the look of every generated document.
"""

import re
import sys

import yaml
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


STYLE = {
    'font_name': 'Calibri',
    'body_size': Pt(10.5),
    'name_size': Pt(20),
    'tagline_size': Pt(10.5),
    'heading_size': Pt(12),
    'margin': Inches(0.6),
    'heading_color': RGBColor(0x1a, 0x1a, 0x1a),
    'muted_color': RGBColor(0x59, 0x59, 0x59),
    'link_color': RGBColor(0x15, 0x63, 0xC0),
    'rule_color': '999999',
}


# ---------------------------------------------------------------------------
# Inline markdown -> runs
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(r'\*\*(.+?)\*\*|\[(.+?)\]\((.+?)\)')


def add_inline_markdown(paragraph, text, base_bold=False):
    """Append text to a paragraph, parsing **bold** and [text](url) spans."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _add_run(paragraph, text[pos:m.start()], bold=base_bold)
        if m.group(1) is not None:
            _add_run(paragraph, m.group(1), bold=True)
        else:
            _add_hyperlink(paragraph, m.group(2), m.group(3))
        pos = m.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:], bold=base_bold)


def _add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.font.name = STYLE['font_name']
    run.font.size = size or STYLE['body_size']
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def _add_hyperlink(paragraph, text, url):
    """Insert a real clickable hyperlink run (python-docx has no native API for this)."""
    part = paragraph.part
    r_id = part.relate_to(
        url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), STYLE['font_name'])
    rPr.append(rFonts)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '%02X%02X%02X' % (STYLE['link_color'][0], STYLE['link_color'][1], STYLE['link_color'][2]))
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(STYLE['body_size'].pt * 2)))
    rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def new_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = STYLE['margin']
    section.bottom_margin = STYLE['margin']
    section.left_margin = STYLE['margin']
    section.right_margin = STYLE['margin']

    normal = doc.styles['Normal']
    normal.font.name = STYLE['font_name']
    normal.font.size = STYLE['body_size']
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    return doc


def _set_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), STYLE['rule_color'])
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Section builders (resume)
# ---------------------------------------------------------------------------

def add_header(doc, name, tagline, contact_parts, immigration_line=None):
    """contact_parts: ordered list of (text, url_or_None) tuples, pipe-separated."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, name, bold=True, size=STYLE['name_size'])

    if tagline:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        _add_run(p2, tagline, italic=True, size=STYLE['tagline_size'], color=STYLE['muted_color'])

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    for i, (text, url) in enumerate(contact_parts):
        if i > 0:
            _add_run(p3, '  |  ', color=STYLE['muted_color'])
        if url:
            _add_hyperlink(p3, text, url)
        else:
            _add_run(p3, text)

    if immigration_line:
        p4 = doc.add_paragraph()
        p4.paragraph_format.space_after = Pt(8)
        _add_run(p4, immigration_line, italic=True, color=STYLE['muted_color'])


def add_section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    _add_run(p, title, bold=True, size=STYLE['heading_size'], color=STYLE['heading_color'])
    _set_bottom_border(p)


def add_summary(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_inline_markdown(p, text)


def add_paragraph_position(doc, theme, dates, role_institution, body_markdown):
    header = doc.add_paragraph()
    header.paragraph_format.space_after = Pt(0)
    tab_stops = header.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.3 - 2 * STYLE['margin'].inches), WD_TAB_ALIGNMENT.RIGHT)
    _add_run(header, theme, bold=True)
    header.add_run('\t')
    _add_run(header, dates, color=STYLE['muted_color'])

    if role_institution:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(2)
        _add_run(sub, role_institution, italic=True)

    body = doc.add_paragraph()
    body.paragraph_format.space_after = Pt(8)
    add_inline_markdown(body, body_markdown)


def add_flat_bullets(doc, items):
    """items: list of (bold_label_or_None, text_markdown) tuples."""
    for label, text in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        if label:
            _add_run(p, f'{label}: ', bold=True)
        add_inline_markdown(p, text)


def add_education(doc, entries):
    for e in entries:
        header = doc.add_paragraph()
        header.paragraph_format.space_after = Pt(0)
        tab_stops = header.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(7.3 - 2 * STYLE['margin'].inches), WD_TAB_ALIGNMENT.RIGHT)
        _add_run(header, e['degree'], bold=True)
        header.add_run('\t')
        _add_run(header, e['dates'], color=STYLE['muted_color'])

        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(6)
        _add_run(sub, f"{e['institution']}, {e['location']}")
        if e.get('gpa'):
            _add_run(sub, f"    GPA: {e['gpa']}")


def add_publications(doc, pubs, scholar_url=None, paper_count=None, citation_count=None):
    heading_extra = ''
    if paper_count and citation_count:
        heading_extra = f' (Google Scholar: {paper_count} papers | {citation_count}+ citations)'
    add_section_heading(doc, f'Selected Publications{heading_extra}')
    for pub in pubs[:5]:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        add_inline_markdown(p, pub)


# ---------------------------------------------------------------------------
# Resume assembly
# ---------------------------------------------------------------------------

def build_resume(content, output_path):
    doc = new_document()
    h = content['header']

    contact_parts = [(h['location'], None), (h['phone'], None), (h['email'], f"mailto:{h['email']}")]
    for key in ('linkedin', 'github'):
        if h.get(key) and h[key].get('url'):
            contact_parts.append((h[key]['label'], h[key]['url']))
    if h.get('languages'):
        contact_parts.append((h['languages'], None))

    add_header(doc, h['name'], h.get('tagline'), contact_parts, h.get('immigration_line'))

    add_section_heading(doc, 'Summary')
    add_summary(doc, content['summary'])

    add_section_heading(doc, 'Work Experience')
    for pos in content['work_experience']:
        add_paragraph_position(doc, pos['theme'], pos['dates'], pos.get('role_institution'), pos['paragraph'])

    add_section_heading(doc, 'Technical Skills')
    add_flat_bullets(doc, [(s['label'], s['text']) for s in content['technical_skills']])

    add_section_heading(doc, 'Work Samples')
    ws_items = []
    for group in content['work_samples']:
        links_md = ', '.join(f"[{i['text']}]({i['url']})" for i in group['items'])
        ws_items.append((group['category'], links_md))
    add_flat_bullets(doc, ws_items)

    add_section_heading(doc, 'Education')
    add_education(doc, content['education'])

    pubs = content['publications']
    add_publications(doc, pubs['items'], pubs.get('scholar_url'), pubs.get('paper_count'), pubs.get('citation_count'))

    add_section_heading(doc, 'Cross-Functional Leadership and Stakeholder Engagement')
    add_flat_bullets(doc, [(c['label'], c['text']) for c in content['cross_functional_leadership']])

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Cover letter assembly
# ---------------------------------------------------------------------------

def build_cover_letter(content, output_path):
    doc = new_document()
    h = content['header']

    p = doc.add_paragraph()
    _add_run(p, h['name'], bold=True, size=STYLE['name_size'])
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    contact_bits = [h.get('location'), h.get('phone'), h.get('email')]
    _add_run(p2, '  |  '.join(b for b in contact_bits if b), color=STYLE['muted_color'])

    recipient = content.get('recipient')
    if recipient:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_after = Pt(10)
        for line in recipient:
            _add_run(rp, line)
            rp.add_run().add_break()

    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_after = Pt(10)
    _add_run(date_p, content.get('date', ''))

    open_p = doc.add_paragraph()
    open_p.paragraph_format.space_after = Pt(8)
    _add_run(open_p, content.get('salutation', 'Dear Hiring Committee,'))

    for para in content['body_paragraphs']:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_inline_markdown(p, para)

    close_p = doc.add_paragraph()
    close_p.paragraph_format.space_before = Pt(8)
    _add_run(close_p, content.get('closing', 'Sincerely,'))
    close_p.add_run().add_break()
    for line in content.get('signature', []):
        _add_run(close_p, line)
        close_p.add_run().add_break()

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) != 5 or sys.argv[1] != 'render' or sys.argv[2] not in ('resume', 'cover_letter'):
        print(__doc__)
        sys.exit(1)

    _, _, doc_type, content_path, output_path = sys.argv
    with open(content_path) as f:
        content = yaml.safe_load(f)

    if doc_type == 'resume':
        build_resume(content, output_path)
    else:
        build_cover_letter(content, output_path)

    print(f'Wrote {output_path}')


if __name__ == '__main__':
    main()
